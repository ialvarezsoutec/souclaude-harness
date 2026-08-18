#!/usr/bin/env python3
"""
Soutec — Convierte un archivo Markdown (.md) en un DOCX corporativo Soutec.

Hermano de `md_to_pdf.py`: comparte con él el núcleo `soutec_md.py` (front-matter,
parseo a bloques y numeración jerárquica), de modo que el Word y el PDF del mismo
`.md` tienen exactamente la misma estructura, numeración y textos. Lo que cambia
es el motor de salida: aquí es **python-docx** (Python puro, sin Word instalado).

    pip install python-docx pillow markdown

Uso:
    python md_to_docx.py ENTRADA.md [SALIDA.docx] [opciones]

Opciones (sobre-escriben el front-matter y los defaults):
    --title / --header / --subtitle / --date / --author / --url / --client-logo
    --no-cover | --no-toc | --no-backcover

Notas de fidelidad respecto al PDF:
  * Las franjas anguladas (banner corrido y de sección) se dibujan como PNG
    generados al vuelo con Pillow y se anclan DETRÁS del texto, así el título
    sigue siendo texto real, editable y visible para el índice de Word.
  * El índice es un campo TOC real: Word lo rellena al abrir (o con F9), por lo
    que los números de página siempre quedan correctos aunque se edite el texto.
"""
import argparse
import datetime
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import soutec_md as core
from soutec_md import (
    parse_front_matter, truthy, extract_title, md_to_blocks, today_es,
    inline_runs, plain,
)

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.normpath(os.path.join(HERE, "..", "assets"))
LOGO = os.path.join(ASSETS, "soutec_logo.png")
ISOTIPO_3D = os.path.join(ASSETS, "soutec_isotipo_3d_color.png")
ISOTIPO_3D_WHITE = os.path.join(ASSETS, "soutec_isotipo_3d_white.png")

# ---------- Geometría (cm), idéntica a la del PDF ----------
PAGE_W_CM, PAGE_H_CM = 21.59, 27.94       # Letter
MARGIN_CM = 2.0
CONTENT_W_CM = PAGE_W_CM - 2 * MARGIN_CM

BANNER_H_CM = 0.84                        # alto de la franja del encabezado
BANNER_SKEW_CM = 0.48
SECBAR_H_CM = 0.88                        # alto de la franja de sección
SECBAR_SKEW_CM = 0.46
SEC_NUM_X_CM = 0.60                       # x del número, desde el borde del papel
SEC_TITLE_X_CM = 1.68                     # x del título, desde el borde del papel

IND_H2, IND_H3, IND_H4 = 0.81, 1.75, 3.00   # sangrías de subtítulos
IND_BODY_FIRST = 0.53                       # sangría de primera línea del cuerpo
IND_LIST_TEXT, IND_LIST_BULLET = 1.65, 1.30

CLIENT_LOGO_H_CM = 2.2                    # logo de cliente en portada
PAGE_LOGO_H_CM = 1.05                     # logo en el encabezado de página

# Fuente del documento. La plantilla corporativa de Word usa Calibri; el manual
# de marca admite Quattrocento Sans para documentos. Se deja configurable aquí
# porque una fuente no instalada en la máquina del lector se sustituye sola.
FONT_BODY = "Calibri"
FONT_MONO = "Consolas"

RENDER_DPI = 200                          # resolución de las franjas rasterizadas


def _rgb(hexval):
    return RGBColor.from_string(hexval.lstrip("#").upper())


# ============================================================
#  Utilidades de XML (lo que python-docx no expone)
# ============================================================
def _shade(el, hexcolor):
    """Sombreado de fondo para un párrafo o una celda."""
    pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor.lstrip("#").upper())
    pr.append(shd)


def _field(paragraph, instr, cached=""):
    """Inserta un campo de Word (PAGE, NUMPAGES, TOC…) con resultado cacheado."""
    r1 = paragraph.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    fc.set(qn("w:dirty"), "true")
    r1._r.append(fc)

    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r2._r.append(it)

    r3 = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)

    r4 = paragraph.add_run(cached)

    r5 = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return [r1, r2, r3, r4, r5]


def _update_fields_on_open(doc):
    """Pide a Word que refresque los campos (el índice) al abrir el archivo."""
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def _anchor_behind(paragraph, image_path, w_cm, h_cm, x_cm, y_cm,
                   rel_h="page", rel_v="paragraph"):
    """Ancla una imagen DETRÁS del texto en una posición absoluta.

    Es la pieza que permite que las franjas anguladas sangren hasta el borde del
    papel sin convertir el título en imagen: el color va detrás, el texto queda
    encima y sigue siendo texto real.
    """
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(w_cm), height=Cm(h_cm))
    inline = run._r.find(qn("w:drawing"))[0]

    anchor = OxmlElement("wp:anchor")
    for k, v in (("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
                 ("simplePos", "0"), ("relativeHeight", "1"), ("behindDoc", "1"),
                 ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1")):
        anchor.set(k, v)

    sp = OxmlElement("wp:simplePos")
    sp.set("x", "0")
    sp.set("y", "0")
    anchor.append(sp)

    ph = OxmlElement("wp:positionH")
    ph.set("relativeFrom", rel_h)
    off = OxmlElement("wp:posOffset")
    off.text = str(int(Cm(x_cm)))
    ph.append(off)
    anchor.append(ph)

    pv = OxmlElement("wp:positionV")
    pv.set("relativeFrom", rel_v)
    offv = OxmlElement("wp:posOffset")
    offv.text = str(int(Cm(y_cm)))
    pv.append(offv)
    anchor.append(pv)

    for tag in ("wp:extent", "wp:effectExtent", "wp:docPr",
                "{http://schemas.openxmlformats.org/drawingml/2006/picture}pic",
                "a:graphic"):
        pass
    # Reutiliza los hijos del inline original (extent, docPr, graphic).
    for child in list(inline):
        inline.remove(child)
        anchor.append(child)
    # wrapNone debe ir antes de docPr/graphic según el esquema.
    wrap = OxmlElement("wp:wrapNone")
    docpr = anchor.find(qn("wp:docPr"))
    anchor.insert(list(anchor).index(docpr), wrap)

    inline.getparent().replace(inline, anchor)
    return run


def _text_w_cm(text, pt_size, bold=True):
    """Ancho aproximado del texto en cm (métricas Helvetica como proxy del
    tipo de Word; se usa solo para dimensionar franjas, con holgura)."""
    from reportlab.pdfbase.pdfmetrics import stringWidth
    font = "Helvetica-Bold" if bold else "Helvetica"
    return stringWidth(text, font, pt_size) / 72.0 * 2.54 * 1.04


# ============================================================
#  Franjas anguladas rasterizadas (PNG transparente)
# ============================================================
def _bar_png(path, w_cm, h_cm, skew_cm, hexcolor):
    """Dibuja el polígono de la franja (rectángulo + punta) sobre transparencia."""
    from PIL import Image, ImageDraw
    px = lambda c: max(1, int(round(c / 2.54 * RENDER_DPI)))
    W, H, S = px(w_cm), px(h_cm), px(skew_cm)
    im = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    d = ImageDraw.Draw(im)
    rgb = tuple(int(hexcolor.lstrip("#")[i:i + 2], 16) for i in (0, 2, 4))
    # (0,0) arriba-izq. La franja es más ancha arriba: la punta baja hacia la der.
    d.polygon([(0, 0), (W, 0), (W - S, H), (0, H)], fill=rgb + (255,))
    im.save(path)
    return path


# ============================================================
#  Estilos del documento
# ============================================================
def _base_styles(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT_BODY
    st.font.size = Pt(11)
    st.font.color.rgb = _rgb(core.BODY_TXT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.35
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(IND_BODY_FIRST)

    def heading(name, size, color, indent, space_before, space_after):
        s = doc.styles[name]
        s.font.name = FONT_BODY
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.italic = False      # Heading 4 de Word viene en cursiva
        s.font.color.rgb = _rgb(color)
        s.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        p = s.paragraph_format
        p.left_indent = Cm(indent)
        p.first_line_indent = Cm(0)
        p.space_before = Pt(space_before)
        p.space_after = Pt(space_after)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.keep_with_next = True
        p.line_spacing = 1.0
        return s

    # Heading 1 es la franja azul: texto blanco, sangría negativa para que el
    # título arranque a SEC_TITLE_X_CM del borde del papel.
    h1 = heading("Heading 1", 13.5, "#FFFFFF",
                 SEC_TITLE_X_CM - MARGIN_CM, 14, 8)
    h1.paragraph_format.space_before = Pt(16)
    heading("Heading 2", 13, core.BLUE, IND_H2, 12, 5)
    heading("Heading 3", 11.5, core.CYAN, IND_H3, 9, 3)
    heading("Heading 4", 10.5, core.CYAN, IND_H4, 7, 2)


def _p(doc_or_cell, text_runs=None, style=None, **fmt):
    p = doc_or_cell.add_paragraph(style=style)
    if text_runs:
        _add_runs(p, text_runs)
    pf = p.paragraph_format
    for k, v in fmt.items():
        setattr(pf, k, v)
    return p


def _add_runs(paragraph, runs, color=None, size=None, bold=None):
    for r in runs:
        run = paragraph.add_run(r["text"])
        run.bold = r["bold"] if bold is None else bold
        run.italic = r["italic"]
        if r["code"]:
            run.font.name = FONT_MONO
            run.font.color.rgb = _rgb(core.DEEP)
        elif r["href"]:
            run.font.color.rgb = _rgb(core.BLUE)
            run.underline = True
        elif color:
            run.font.color.rgb = _rgb(color)
        if size:
            run.font.size = Pt(size)
    return paragraph


# ============================================================
#  Encabezado, pie y franjas
# ============================================================
def _build_header(section, header_txt, client_logo, tmpdir):
    """Franja azul angulada + logo (cliente o Soutec) a la derecha."""
    hdr = section.header
    hdr.is_linked_to_previous = False
    p = hdr.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    bw = max(6.8, MARGIN_CM + _text_w_cm(header_txt, 9.5) + 0.6)
    png = _bar_png(os.path.join(tmpdir, "hdr_bar.png"),
                   bw + BANNER_SKEW_CM, BANNER_H_CM, BANNER_SKEW_CM, core.BLUE)
    # La franja arranca en el borde del papel; el texto, en el margen.
    _anchor_behind(p, png, bw + BANNER_SKEW_CM, BANNER_H_CM,
                   x_cm=0, y_cm=-0.12, rel_h="page", rel_v="paragraph")
    run = p.add_run(header_txt)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = _rgb("#FFFFFF")

    # Logo a la derecha: anclado por posición absoluta respecto a la página, no
    # con tabulaciones (las del estilo Header interferirían) ni alineación.
    if client_logo:
        lw, lh = _logo_size(client_logo, PAGE_LOGO_H_CM, max_w_cm=4.6)
        logo_path = client_logo
    else:
        lw = 3.0
        lh = lw * 51.0 / 279.0
        logo_path = LOGO
    _anchor_behind(p, logo_path, lw, lh,
                   x_cm=PAGE_W_CM - MARGIN_CM - lw,
                   y_cm=(BANNER_H_CM - lh) / 2 - 0.12,
                   rel_h="page", rel_v="paragraph")


def _build_footer(section, copyright_txt):
    """Copyright a la izquierda y 'Página X de Y' a la derecha, sin línea.

    Se arma con una tabla de dos celdas sin bordes en lugar de una tabulación:
    los estilos Footer traen tabulaciones propias que capturarían el tabulador
    antes que la nuestra, y su posición varía entre Word y LibreOffice.
    """
    ftr = section.footer
    ftr.is_linked_to_previous = False
    p0 = ftr.paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    for r in list(p0.runs):
        r._r.getparent().remove(r._r)

    t = ftr.add_table(rows=1, cols=2, width=Cm(CONTENT_W_CM))
    t.autofit = False
    t.columns[0].width = Cm(CONTENT_W_CM * 0.62)
    t.columns[1].width = Cm(CONTENT_W_CM * 0.38)
    _no_borders(t)

    def cell_par(cell, align):
        cell.width = cell.width
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.alignment = align
        return p

    left = cell_par(t.rows[0].cells[0], WD_ALIGN_PARAGRAPH.LEFT)
    r = left.add_run(copyright_txt)
    r.font.size = Pt(8.5)
    r.font.color.rgb = _rgb(core.FOOT_GREY)

    right = cell_par(t.rows[0].cells[1], WD_ALIGN_PARAGRAPH.RIGHT)
    runs = [right.add_run("Página ")]
    runs += _field(right, " PAGE ", "1")
    runs.append(right.add_run(" de "))
    runs += _field(right, " NUMPAGES ", "1")
    for run in runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = _rgb(core.FOOT_GREY)


def _no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def _clear_header_footer(section):
    for part in (section.header, section.footer):
        part.is_linked_to_previous = False
        for p in part.paragraphs:
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
        part.paragraphs[0].paragraph_format.space_after = Pt(0)


def _section_banner(doc, num, title, tmpdir, idx):
    """Heading 1 con la franja azul anclada detrás del texto."""
    p = doc.add_paragraph(style="Heading 1")
    pf = p.paragraph_format
    # Sangría francesa: el número cuelga a SEC_NUM_X_CM del borde del papel y
    # el título (incluidas sus líneas siguientes) arranca en SEC_TITLE_X_CM.
    pf.left_indent = Cm(SEC_TITLE_X_CM - MARGIN_CM)
    pf.first_line_indent = Cm(-(SEC_TITLE_X_CM - SEC_NUM_X_CM))

    bw_max = PAGE_W_CM - 0.4 - SECBAR_SKEW_CM
    avail = bw_max - SEC_TITLE_X_CM - 0.7
    tw = _text_w_cm(title, 13.5)
    # Si el título no cabe en una línea, el párrafo hará wrap: la franja tiene
    # que crecer con él (igual que en el PDF) o la segunda línea quedaría fuera.
    nlines = max(1, int(tw / avail) + (1 if tw % avail else 0))
    bar_h = SECBAR_H_CM if nlines == 1 else 0.62 * nlines + 0.30
    bw = min(max(7.4, SEC_TITLE_X_CM + min(tw, avail) + 0.7), bw_max)
    png = _bar_png(os.path.join(tmpdir, f"sec_{idx}.png"),
                   bw + SECBAR_SKEW_CM, bar_h, SECBAR_SKEW_CM, core.BLUE)
    # Ancla respecto a la LÍNEA, no al párrafo: el `space_before` del Heading 1
    # se cuenta dentro del párrafo en algunos renderizadores y desplazaría la
    # franja sobre el texto (se nota tras una tabla).
    _anchor_behind(p, png, bw + SECBAR_SKEW_CM, bar_h,
                   x_cm=0, y_cm=-(bar_h - 0.55 * nlines) / 2,
                   rel_h="page", rel_v="line")

    # El número va a SEC_NUM_X_CM del borde: se coloca con una tabulación
    # negativa respecto al margen (Word admite posiciones negativas).
    pf.tab_stops.add_tab_stop(Cm(SEC_TITLE_X_CM - MARGIN_CM), WD_TAB_ALIGNMENT.LEFT)
    r = p.add_run(str(num))
    r.font.color.rgb = _rgb("#FFFFFF")
    r.bold = True
    p.add_run("\t")
    r = p.add_run(title)
    r.font.color.rgb = _rgb("#FFFFFF")
    r.bold = True
    return p


# ============================================================
#  Tablas, código y listas
# ============================================================
def _add_table(doc, block):
    head = block.get("head") or []
    rows = block.get("rows") or []
    ncol = max([len(head)] + [len(r) for r in rows]) if (head or rows) else 1
    nrow = len(rows) + (1 if head else 0)
    if not nrow:
        return
    t = doc.add_table(rows=nrow, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

    def fill(cell, html, bold=False, white=False, size=9.5):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        runs = inline_runs(html) or [{"text": "", "bold": False, "italic": False,
                                      "code": False, "href": None}]
        for r in runs:
            run = p.add_run(r["text"])
            run.bold = bold or r["bold"]
            run.italic = r["italic"]
            run.font.size = Pt(size)
            if white:
                run.font.color.rgb = _rgb("#FFFFFF")
            elif r["code"]:
                run.font.name = FONT_MONO

    r0 = 0
    if head:
        for j in range(ncol):
            cell = t.rows[0].cells[j]
            fill(cell, head[j] if j < len(head) else "", bold=True, white=True)
            _shade(cell._tc, core.BLUE)
        t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        r0 = 1
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = t.rows[r0 + i].cells[j]
            fill(cell, row[j] if j < len(row) else "")
            if i % 2 == 1:
                _shade(cell._tc, core.ZEBRA)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _add_code(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.4)
    pf.right_indent = Cm(0.2)
    pf.line_spacing = 1.0
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _shade(p._p, core.CODEBG)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = FONT_MONO
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(core.CARBON)
    return p


def _add_list_item(doc, html, ordered, n):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(IND_LIST_TEXT)
    pf.first_line_indent = Cm(IND_LIST_BULLET - IND_LIST_TEXT)
    pf.space_after = Pt(3)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.tab_stops.add_tab_stop(Cm(IND_LIST_TEXT), WD_TAB_ALIGNMENT.LEFT)
    mark = f"{n}." if ordered else "•"
    r = p.add_run(mark)
    r.bold = True
    r.font.color.rgb = _rgb(core.BLUE)
    p.add_run("\t")
    _add_runs(p, inline_runs(html))
    return p


# ============================================================
#  Portada, índice y contraportada
# ============================================================
def _logo_size(path, height_cm, max_w_cm=9.0):
    try:
        from PIL import Image
        with Image.open(path) as im:
            ar = im.width / float(im.height)
    except Exception:
        ar = 1.0
    w = height_cm * ar
    if w > max_w_cm:
        return max_w_cm, max_w_cm / ar
    return w, height_cm


def _cover(doc, title, subtitle, date_str, author, client_logo, tmpdir):
    """Portada: fecha arriba a la derecha, isotipo 3D sangrando por la izquierda
    y bloque logo/título/subtítulo a la derecha, como en la plantilla."""
    from PIL import Image

    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.space_after = Pt(0)
    r = p.add_run(date_str)
    r.font.size = Pt(9)
    r.font.color.rgb = _rgb(core.FOOT_GREY)

    # Isotipo: mismas medidas que el PDF (alto 18.4 cm, borde derecho al margen).
    # Como en Word el desplazamiento respecto a la página no puede ser negativo,
    # se recorta la parte que sangra fuera del papel y se ancla el resto en x=0.
    iso_h = 18.4
    with Image.open(ISOTIPO_3D) as im:
        ar = im.width / float(im.height)
        iso_w = iso_h * ar
        x_left = PAGE_W_CM - MARGIN_CM - iso_w      # negativo: sangra
        if x_left < 0:
            frac = -x_left / iso_w
            crop = im.crop((int(im.width * frac), 0, im.width, im.height))
            iso_w += x_left
            x_left = 0
        else:
            crop = im.copy()
        cut = os.path.join(tmpdir, "cover_iso.png")
        crop.save(cut)
    # El tope del isotipo va a 3.1 cm del borde superior del papel (igual que en
    # el PDF), anclado respecto a la página para no depender del flujo de texto.
    _anchor_behind(p, cut, iso_w, iso_h, x_cm=x_left, y_cm=3.1,
                   rel_h="page", rel_v="page")

    # Hueco del logo (cliente o Soutec)
    lp = doc.add_paragraph()
    lpf = lp.paragraph_format
    lpf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lpf.first_line_indent = Cm(0)
    lpf.space_before = Pt(Cm(PAGE_H_CM - 6.45 - MARGIN_CM - 1.4).pt)
    lpf.space_after = Pt(10)
    run = lp.add_run()
    if client_logo:
        w, h = _logo_size(client_logo, CLIENT_LOGO_H_CM)
        run.add_picture(client_logo, height=Cm(h))
    else:
        run.add_picture(LOGO, width=Cm(4.6))

    tp = doc.add_paragraph()
    tpf = tp.paragraph_format
    tpf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tpf.first_line_indent = Cm(0)
    tpf.space_after = Pt(2)
    tpf.line_spacing = 1.15
    r = tp.add_run(title)
    r.bold = True
    r.font.size = Pt(17.5)
    r.font.color.rgb = _rgb(core.COVER_GREY)

    if subtitle:
        sp = doc.add_paragraph()
        spf = sp.paragraph_format
        spf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        spf.first_line_indent = Cm(0)
        spf.space_after = Pt(2)
        r = sp.add_run(subtitle)
        r.bold = True
        r.font.size = Pt(17.5)
        r.font.color.rgb = _rgb(core.COVER_GREY)

    if author:
        ap = doc.add_paragraph()
        apf = ap.paragraph_format
        apf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        apf.first_line_indent = Cm(0)
        r = ap.add_run(author)
        r.font.size = Pt(10)
        r.font.color.rgb = _rgb(core.FOOT_GREY)


def _toc(doc):
    """Título 'Contenidos' + campo TOC real (Word rellena los números)."""
    h = doc.add_paragraph()
    hpf = h.paragraph_format
    hpf.first_line_indent = Cm(0)
    hpf.space_after = Pt(10)
    hpf.keep_with_next = True
    r = h.add_run("Contenidos")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = _rgb(core.BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    _field(p, ' TOC \\o "1-3" \\h \\z \\u ',
           "Pulsa Ctrl+E, F9 (o clic derecho ▸ Actualizar campos) para generar el índice.")


def _backcover(doc, url, tmpdir):
    """Contraportada: página azul completa con el isotipo 3D blanco y la URL."""
    from PIL import Image

    px = lambda c: int(round(c / 2.54 * RENDER_DPI))
    W, H = px(PAGE_W_CM), px(PAGE_H_CM)
    bg = Image.new("RGB", (W, H), tuple(
        int(core.BACK_BLUE.lstrip("#")[i:i + 2], 16) for i in (0, 2, 4)))
    with Image.open(ISOTIPO_3D_WHITE) as iso:
        iso = iso.convert("RGBA")
        ih = px(19.0)
        iw = int(iso.width * ih / iso.height)
        iso = iso.resize((iw, ih), Image.LANCZOS)
        # mismas coordenadas que el PDF: x=-8.5 cm, tope a 3.2 cm sobre el papel
        bg.paste(iso, (px(-8.5), px(-3.2)), iso)
    full = os.path.join(tmpdir, "backcover.png")
    bg.save(full)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    # Se dibuja con un pelo de sangrado (2.5 mm) para que el redondeo de
    # renderizado no deje un filete blanco en los bordes derecho e inferior.
    _anchor_behind(p, full, PAGE_W_CM + 0.25, PAGE_H_CM + 0.25, x_cm=0, y_cm=0,
                   rel_h="page", rel_v="page")

    u = doc.add_paragraph()
    upf = u.paragraph_format
    upf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    upf.first_line_indent = Cm(0)
    upf.space_before = Pt(Cm(PAGE_H_CM - 2.8 - MARGIN_CM - 1.0).pt)
    r = u.add_run(url)
    r.font.size = Pt(11.5)
    r.font.color.rgb = _rgb("#FFFFFF")


def _setup_section(section, margins=True):
    section.page_width = Cm(PAGE_W_CM)
    section.page_height = Cm(PAGE_H_CM)
    m = Cm(MARGIN_CM) if margins else Cm(0)
    section.left_margin = section.right_margin = Cm(MARGIN_CM)
    section.top_margin = Cm(MARGIN_CM + 0.85)
    section.bottom_margin = Cm(MARGIN_CM)
    section.header_distance = Cm(1.35)
    section.footer_distance = Cm(1.05)
    return section


# ============================================================
#  Render principal
# ============================================================
def render(md_path, out_path, args):
    import tempfile

    with open(md_path, "r", encoding="utf-8") as f:
        raw = f.read()
    meta, body = parse_front_matter(raw)

    if args.title:
        title = args.title
    else:
        title, body = extract_title(body, meta)
    header_txt = args.header or meta.get("header") or title
    subtitle = args.subtitle or meta.get("subtitle", "")
    author = args.author or meta.get("author", "")
    url = args.url or meta.get("url", "www.soutec-group.com")

    client_logo = args.client_logo or meta.get("client_logo") or ""
    if client_logo:
        client_logo = os.path.expanduser(client_logo)
        if not os.path.isabs(client_logo):
            client_logo = os.path.join(os.path.dirname(os.path.abspath(md_path)),
                                       client_logo)
        if not os.path.isfile(client_logo):
            sys.exit(f"No se encontró el logo del cliente: {client_logo}\n"
                     "Envía la imagen del logo (PNG/JPG, preferible con fondo "
                     "transparente) o corrige la ruta con --client-logo.")

    want_cover = not args.no_cover
    want_toc = not args.no_toc
    want_backcover = not args.no_backcover and not (
        "backcover" in meta and not truthy(meta.get("backcover", "true")))

    date_str = args.date or meta.get("date", today_es())
    year = datetime.date.today().year
    copyright_txt = meta.get("copyright") or \
        f"©{year} Soutec – Todos los Derechos Reservados"

    blocks = md_to_blocks(body)
    has_headings = any(b["t"] in ("section", "h") for b in blocks)

    tmpdir = tempfile.mkdtemp(prefix="soutec_docx_")
    doc = Document()
    _base_styles(doc)
    _setup_section(doc.sections[0])

    if want_cover:
        _clear_header_footer(doc.sections[0])
        doc.sections[0].top_margin = Cm(MARGIN_CM)
        _cover(doc, title, subtitle, date_str, author, client_logo or None, tmpdir)
        body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        _setup_section(body_sec)
    else:
        body_sec = doc.sections[0]
    _build_header(body_sec, header_txt, client_logo or None, tmpdir)
    _build_footer(body_sec, copyright_txt)

    if want_toc and has_headings:
        _toc(doc)
        doc.add_page_break()

    secidx = 0
    for b in blocks:
        t = b["t"]
        if t == "section":
            secidx += 1
            _section_banner(doc, b["num"], plain(b["title"]), tmpdir, secidx)
        elif t == "h":
            lvl = min(max(b["level"], 2), 4)
            p = doc.add_paragraph(style=f"Heading {lvl}")
            num = b.get("num") or ""
            if num:
                r = p.add_run(num + "   ")
                r.bold = True
            _add_runs(p, inline_runs(b["text"]), bold=True)
        elif t == "p":
            if b["text"]:
                _p(doc, inline_runs(b["text"]))
        elif t == "callout":
            _p(doc, inline_runs(b["text"]))
        elif t == "list":
            for i, it in enumerate(b["items"], 1):
                _add_list_item(doc, it, b["ordered"], i)
        elif t == "table":
            _add_table(doc, b)
        elif t == "code":
            _add_code(doc, b["text"])
        elif t == "hr":
            hp = doc.add_paragraph()
            pr = hp._p.get_or_add_pPr()
            bd = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), core.RULE.lstrip("#"))
            bd.append(bot)
            pr.append(bd)

    if want_backcover:
        back_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        _setup_section(back_sec)
        back_sec.top_margin = Cm(0)
        back_sec.bottom_margin = Cm(0)
        back_sec.left_margin = back_sec.right_margin = Cm(0)
        _clear_header_footer(back_sec)
        _backcover(doc, url, tmpdir)

    _update_fields_on_open(doc)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(
        description="Convierte Markdown a DOCX corporativo Soutec (python-docx).")
    ap.add_argument("input")
    ap.add_argument("output", nargs="?")
    ap.add_argument("--title")
    ap.add_argument("--header")
    ap.add_argument("--subtitle")
    ap.add_argument("--date")
    ap.add_argument("--author")
    ap.add_argument("--url")
    ap.add_argument("--client-logo", help="Ruta al logo del cliente para la portada "
                    "y el encabezado. Si se omite, se usa el logo de Soutec.")
    ap.add_argument("--no-cover", action="store_true")
    ap.add_argument("--no-toc", action="store_true")
    ap.add_argument("--no-backcover", action="store_true")
    args = ap.parse_args()

    if not os.path.isfile(args.input):
        sys.exit(f"No existe el archivo: {args.input}")
    out = args.output or os.path.splitext(args.input)[0] + ".docx"
    render(args.input, out, args)
    print(f"DOCX generado: {out}")


if __name__ == "__main__":
    main()
